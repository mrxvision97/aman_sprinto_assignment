"""
Document parsing service.
Primary: Unstructured.io API (handles complex PDFs, multi-column, image-based).
Fallback: pdfplumber (PDF) / python-docx (DOCX).
"""
import re
import io
import asyncio
import mimetypes
from typing import Optional
import pdfplumber
from docx import Document


SECTION_PATTERNS = re.compile(
    r"^(EXPERIENCE|WORK EXPERIENCE|EMPLOYMENT|EDUCATION|SKILLS|TECHNICAL SKILLS|"
    r"PROJECTS|CERTIFICATIONS|SUMMARY|OBJECTIVE|PROFILE|CONTACT|PUBLICATIONS|"
    r"AWARDS|LANGUAGES|INTERESTS|VOLUNTEER|REFERENCES|ACHIEVEMENTS|EXPERTISE)S?\s*$",
    re.IGNORECASE | re.MULTILINE
)


async def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Parse resume using Unstructured.io API first.
    Falls back to pdfplumber/python-docx if API fails or returns too little text.
    """
    # Try Unstructured.io first
    result = await _parse_with_unstructured(file_bytes, filename)
    if result["parsed_text"] and len(result["parsed_text"].strip()) >= 200:
        return result

    # Fallback to local parsers
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        fallback = _parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        fallback = _parse_docx(file_bytes)
    else:
        fallback = {
            "parsed_text": None,
            "parse_report": {
                "confidence": 0,
                "method": "unsupported",
                "sections_found": [],
                "warnings": [f"Unsupported file type: {ext}"],
                "char_count": 0,
            },
        }

    # If fallback got more text, use it; otherwise return whichever has something
    fallback_len = len((fallback.get("parsed_text") or "").strip())
    primary_len = len((result.get("parsed_text") or "").strip())
    return fallback if fallback_len >= primary_len else result


async def _parse_with_unstructured(file_bytes: bytes, filename: str) -> dict:
    """Call Unstructured.io Serverless API."""
    try:
        import unstructured_client
        from unstructured_client.models import operations, shared
        from unstructured_client.utils import RetryConfig, BackoffStrategy
        from app.config import get_settings

        api_key = get_settings().unstructured_api_key
        if not api_key:
            return {"parsed_text": None, "parse_report": {"confidence": 0, "method": "unstructured_api", "sections_found": [], "warnings": ["No Unstructured API key"], "char_count": 0}}

        client = unstructured_client.UnstructuredClient(
            api_key_auth=api_key,
            server_url="https://api.unstructuredapp.io",
            retry_config=RetryConfig(
                strategy="backoff",
                retry_connection_errors=True,
                backoff=BackoffStrategy(
                    initial_interval=500,
                    max_interval=60000,
                    exponent=1.5,
                    max_elapsed_time=900000,  # 15 minutes
                ),
            ),
        )

        ext = filename.lower().split(".")[-1]
        is_pdf = ext == "pdf"
        is_image = ext in ("png", "jpg", "jpeg", "tiff", "bmp", "heic")

        # Use stronger layout-aware parsing for PDFs/images.
        strategy = "hi_res" if (is_pdf or is_image) else "auto"

        # MIME hint can help Unstructured with edge cases.
        content_type, _ = mimetypes.guess_type(filename)
        content_type = content_type or None

        def _call():
            files = shared.Files(content=file_bytes, file_name=filename)
            params = shared.PartitionParameters(
                files=files,
                strategy=strategy,
                languages=["eng"],
                include_page_breaks=True,
                unique_element_ids=True,
                coordinates=False,
                content_type=content_type,
            )

            # Enable more robust handling for PDFs (parallel page splitting + table extraction).
            if is_pdf:
                params.split_pdf_page = True
                params.split_pdf_allow_failed = True
                params.split_pdf_concurrency_level = 10
                # Enable table extraction for PDFs by default (skip list empty).
                params.skip_infer_table_types = []

            req = operations.PartitionRequest(partition_parameters=params)
            resp = client.general.partition(request=req)
            return resp.elements, strategy, content_type

        elements, strategy_used, content_type_used = await asyncio.to_thread(_call)

        lines = []
        for el in (elements or []):
            el_type = el.get("type", "") if isinstance(el, dict) else getattr(el, "type", "")
            text = (el.get("text", "") if isinstance(el, dict) else getattr(el, "text", "")).strip()
            if not text:
                continue
            if el_type in ("Title", "Header"):
                lines.append(f"\n{text.upper()}\n")
            else:
                lines.append(text)

        full_text = "\n".join(lines)
        char_count = len(full_text.strip())
        sections = _detect_sections(full_text)

        return {
            "parsed_text": _clean_text(full_text) if char_count >= 100 else None,
            "parse_report": {
                "confidence": min(100, 70 + len(sections) * 5),
                "method": "unstructured_api",
                "strategy": strategy_used,
                "content_type": content_type_used,
                "sections_found": sections,
                "warnings": [] if char_count >= 100 else ["Unstructured returned very little text"],
                "char_count": char_count,
                "elements_count": len(elements or []),
            },
        }
    except Exception as e:
        return {
            "parsed_text": None,
            "parse_report": {
                "confidence": 0,
                "method": "unstructured_api",
                "sections_found": [],
                "warnings": [f"Unstructured.io API failed: {str(e)}"],
                "char_count": 0,
            },
        }


def _parse_pdf(file_bytes: bytes) -> dict:
    """Fallback: pdfplumber PDF extraction."""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if text:
                    pages_text.append(text)

            full_text = "\n".join(pages_text)
            char_count = len(full_text.strip())

            if char_count < 100:
                return {
                    "parsed_text": None,
                    "parse_report": {
                        "confidence": 0,
                        "method": "pdf_fallback",
                        "sections_found": [],
                        "warnings": ["Very little text — file may be image-based or corrupted"],
                        "char_count": char_count,
                    },
                }

            sections = _detect_sections(full_text)
            confidence = min(100, 50 + len(sections) * 10 + min(char_count // 100, 30))
            return {
                "parsed_text": _clean_text(full_text),
                "parse_report": {
                    "confidence": confidence,
                    "method": "pdf_fallback",
                    "sections_found": sections,
                    "warnings": [],
                    "char_count": char_count,
                    "page_count": len(pdf.pages),
                },
            }
    except Exception as e:
        return {
            "parsed_text": None,
            "parse_report": {
                "confidence": 0,
                "method": "pdf_fallback",
                "sections_found": [],
                "warnings": [f"PDF fallback failed: {str(e)}"],
                "char_count": 0,
            },
        }


def _parse_docx(file_bytes: bytes) -> dict:
    """Fallback: python-docx extraction."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                lines.append(f"\n{text.upper()}\n")
            else:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))

        full_text = "\n".join(lines)
        char_count = len(full_text.strip())
        sections = _detect_sections(full_text)
        confidence = min(100, 60 + len(sections) * 10 + min(char_count // 100, 20))
        return {
            "parsed_text": _clean_text(full_text),
            "parse_report": {
                "confidence": confidence,
                "method": "docx_fallback",
                "sections_found": sections,
                "warnings": [],
                "char_count": char_count,
            },
        }
    except Exception as e:
        return {
            "parsed_text": None,
            "parse_report": {
                "confidence": 0,
                "method": "docx_fallback",
                "sections_found": [],
                "warnings": [f"DOCX fallback failed: {str(e)}"],
                "char_count": 0,
            },
        }


def _detect_sections(text: str) -> list[str]:
    found = set()
    for match in SECTION_PATTERNS.finditer(text):
        found.add(match.group(0).strip().upper())
    return sorted(found)


def _clean_text(text: str) -> str:
    text = re.sub(r"[●■►▪•]", "-", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()
